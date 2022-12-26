import sqlite3
import sys
import docx
import os
import base64
from docx.shared import Pt, RGBColor
import xlrd
import datetime
import math

conn = sqlite3.connect('orders30.db')
cur = conn.cursor()


def convert_File(pathToFile):
    filename = pathToFile
    with open(filename, 'rb') as file:
        photo = file.read()
    encoded = base64.b64encode(photo)
    return encoded


def printFormatRecod(records, header):
    dictMaxRow = dict()
    for i in range(len(records)):
        for j in range(len(header)):
            value = str(records[i][j])
            if j not in dictMaxRow:
                dictMaxRow[j] = max(len(value)+1, len(header[j])+1)
            else:
                dictMaxRow[j] = max(dictMaxRow[j], len(value)+1)
    autoForm = ""
    for i in dictMaxRow.keys():
        if autoForm!="":
            autoForm=autoForm+"| "
        autoForm = autoForm+"{0["+str(i)+"]:<"+str(dictMaxRow[i])+"}"
    print(autoForm.format(header))
    for i in records:
        print(autoForm.format(i))
    print('')



#complex request
def getDocResume(tupleData):
    doc = docx.Document()
    doc.add_heading('Інформація про лікаря', 0)
    records = ['ПІБ',
                'Фото',
                'Дата народження',
                'Стать',
                'Відділення',
                'Посада',
                'Заробітна плата'
    ]
    menuTable = doc.add_table(rows=1, cols=2)
    menuTable.style='Table Grid'
    font = menuTable.style.font
    font.name ='Times New Roman' #Стиль шрифта
    font.size = Pt(16) #Размер шрифта
    hdr_Cells = menuTable.rows[0].cells
    hdr_Cells[0].text = "Поле"
    hdr_Cells[1].text = 'Значення'

    for ID in records:
        row_Cells = menuTable.add_row().cells
        row_Cells[0].text = str(ID)
        if str(ID) == "Фото":
            paragraph =  row_Cells[1].paragraphs[0]
            run = paragraph.add_run()
            decoded = base64.b64decode(tupleData[6])
            f = open("change10.docx", "wb")
            f.write(decoded)
            f.close()
            f = open("change10.docx", "rb")
            run.add_picture(f, width = 4400000, height = 4000000)
            f.close()
        elif str(ID)=='ПІБ':
            row_Cells[1].text = tupleData[1]
        elif str(ID)=='Дата народження':
            row_Cells[1].text = tupleData[4]
        elif str(ID)=='Відділення':
            row_Cells[1].text = tupleData[0]
        elif str(ID)=='Посада':
            row_Cells[1].text = tupleData[2]
        elif str(ID)=='Заробітна плата':
            row_Cells[1].text = str(tupleData[5])
        elif str(ID)=='Стать':
            row_Cells[1].text = str(tupleData[3])

    for cell in menuTable.columns[0].cells:
        cell.width = 8000

    myDocuments = os.path.join(os.path.expanduser('~'),'Downloads')
    print("Document saved by the next path: "+myDocuments+"\doctor.docx")
    doc.save(myDocuments+"\doctor.docx")
    #os.system(myDocuments+"\coctor.docx")


def createDB():
    cur.execute("""CREATE TABLE IF NOT EXISTS hospitalDepartments( 
    ID INTEGER PRIMARY KEY AUTOINCREMENT,
    deparName TEXT NOT NULL,
    shortDescrp TEXT,
    fullDescrp BLOB);""")

    cur.execute("""CREATE TABLE IF NOT EXISTS doctors(
    ID INTEGER PRIMARY KEY AUTOINCREMENT,
    fulName TEXT NOT NULL,
    jobTitle TEXT NOT NULL,
    sex TEXT NOT NULL,
    birthDate DATE,
    Salary INTEGER NOT NULL,
    photo BLOB,
    depertID INTEGER,
    FOREIGN KEY (depertID) REFERENCES hospitalDepartments(ID));""")

    cur.execute("""CREATE TABLE IF NOT EXISTS patients(
    ID INTEGER PRIMARY KEY AUTOINCREMENT,
    fulName TEXT NOT NULL,
    sex TEXT NOT NULL,
    birthDate DATE,
    receiptDate DATE NOT NULL,
    doctorID INTEGER,
    FOREIGN KEY (doctorID) REFERENCES doctors(ID));""")

    cur.execute("""CREATE TABLE IF NOT EXISTS diagnos(
    ID INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL,
    description TEXT,
    recomendation TEXT);""")

    cur.execute("""CREATE TABLE IF NOT EXISTS patients_diagnos(
    patients_id INTEGER NOT NULL,
    diagnos_id INTEGER NOT NULL,
    FOREIGN KEY (patients_id) REFERENCES patients(ID)
    FOREIGN KEY (diagnos_id) REFERENCES diagnos(ID));""")
    conn.commit()



def inputData():
    dataHospitalDepartments = [
    ( 'інтенсивної терапії', 'Коротки опис ... ', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'хірургічне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'травматологічне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'урологічне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'гнійно-септичне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'гінекологічне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'анестезіології ', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'терапевтичне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'неврологічне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'дитяче', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'інфекційне боксоване', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'вагітних', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'ЛОР-відділення', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'офтальмологічне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ( 'приймальне', 'Коротки опис ...', convert_File('\\python_laba\\resourse\\pologennyaProViddil.docx')),
    ]
    dataDoctors = [
    ( 'МЕЛЬНИК МИКОЛА ІВАНОВИЧ', 'Лікар-методист', 'чоловік', '1960-08-16', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 1),
    ( 'КОВАЛЕНКО МИКОЛА ІВАНОВИЧ', 'Лікар-методист', 'чоловік', '1960-09-11', 15000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 1),
    ( 'БОНДАРЕНКО МИКОЛА ІВАНОВИЧ', 'Лікар-методист', 'чоловік', '1993-01-21', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 1),
    ( 'БОЙКО МИКОЛА ІВАНОВИЧ', 'Лікар-методист', 'чоловік', '1967-05-17', 15000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 2),
    ( 'МЕЛЬНИК ВАСИЛЬ ІВАНОВИЧ', 'Лікар-методист', 'чоловік', '1992-02-23', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 1),
    ( 'КОВАЛЕНКО ВОЛОДИМИР ІВАНОВИЧ', 'Лікар-методист', 'чоловік', '1987-02-18', 15000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 1),
    ( 'ШЕВЧЕНКО ОЛЕНА БОРИСІВНА', 'Лікар-терапевт', 'жінка', '1988-10-30', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 2),
    ( 'БОЙКО ГАННА МОКОЛАЇВНА', 'Лікар-терапевт', 'жінка', '1988-01-26', 15000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 3),
    ( 'ШЕВЧЕНКО ВОЛОДИМИР ІВАНОВИЧ', 'Лікар-терапевт', 'чоловік', '1987-11-09', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 2),
    ( 'ОЛІЙНИК МИКОЛА ІВАНОВИЧ', 'Лікар-терапевт', 'чоловік', '1990-08-16', 15000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 3),
    ( 'МЕЛЬНИК ВОЛОДИМИР ІВАНОВИЧ', 'Лікар-терапевт', 'чоловік', '1983-12-03', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 2),
    ('КОВАЛЕНКО ОЛЕКСАНДР МИКОЛАЙОВИЧ','Лікар-терапевт','чоловік','1991-08-16',15000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 3),
    ( 'КОВАЛЕНКО МИКОЛА МИКОЛАЙОВИЧ', 'Лікар-педіатр', 'чоловік', '1979-09-05', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 3),
    ( 'КРАВЧЕНКО АЛІНА ДМИТРІВНА', 'Лікар-педіатр', 'жінка', '1977-08-16', 15000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 2),
    ( 'ТКАЧЕНКО МИКОЛА ІВАНОВИЧ', 'Лікар-педіатр', 'чоловік', '1985-08-04', 12000, convert_File('\\python_laba\\resourse\\doctor.jpg'), 3)
    ]
    dataDiagnoz =  [
    ( 'стенокардіяї', 'Коротки опис ... ', 'Рекомендації по лікуванню ... '), ( 'тахікардія', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),
    ( 'кардіосклероз', 'Коротки опис ...', 'Рекомендації по лікуванню ... '), ( 'бронхіт', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),
    ( 'склеродермія', 'Коротки опис ...', 'Рекомендації по лікуванню ... '), ( 'поліартрит', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),
    ( 'Ентероптоз ', 'Коротки опис ...', 'Рекомендації по лікуванню ... '), ( 'Сколіоз', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),
    ( 'астма', 'Коротки опис ...', 'Рекомендації по лікуванню ... '), ( 'атеросклероз', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),
    ( 'діабет', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),( 'інсульт', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),
    ( 'малярія', 'Коротки опис ...', 'Рекомендації по лікуванню ... '), ( 'інфаркт', 'Коротки опис ...', 'Рекомендації по лікуванню ... '),
    ( 'наркоманія', 'Коротки опис ...','Рекомендації по лікуванню ... '),
    ]
    dataManyToMany = [
    ( 1, 1), ( 1, 2), ( 2, 2),( 2, 3),
    ( 2, 4), ( 3, 4), ( 5, 6), ( 6, 7),
    ( 7, 1), ( 7, 2), ( 8, 2),( 8, 3),
    ( 9, 4), ( 10, 4), ( 11, 6), ( 12, 10),
    ]
    def readPanFromExcel():
        #from https://stackoverflow.com/questions/1108428/how-do-i-read-a-date-in-excel-format-in-python
        def xldate_to_datetime(xldatetime):
            tempDate = datetime.datetime(1899, 12, 31)
            (days, portion) = math.modf(xldatetime)
            deltaDays = datetime.timedelta(days=days)
            secs = int(24 * 60 * 60 * portion)
            detlaSeconds = datetime.timedelta(seconds=secs)
            TheTime = (tempDate + deltaDays + detlaSeconds )
            return TheTime.strftime("%Y-%m-%d")
            #return TheTime.strftime("%Y-%m-%d %H:%M:%S")

        workbook = xlrd.open_workbook("\\python_laba\\resourse\\dataPatients.xls")
        worksheet = workbook.sheet_by_index(0)
        arrTuple = []
        for i in range(0, 15):
            tempAr = []
            for j in range(0, 5):
                if j==4:
                    tempAr.append(int(worksheet.cell_value(i, j)))
                elif j==2 or j==3:
                    strVal = xldate_to_datetime(int(worksheet.cell_value(i, j)))
                    tempAr.append(strVal)
                else:
                    tempAr.append(str(worksheet.cell_value(i, j)).rstrip())
            arrTuple.append(tuple(tempAr))
        return arrTuple
    dataPatient = readPanFromExcel()

    cur.executemany("INSERT INTO hospitalDepartments VALUES(null, ?, ?, ?);", dataHospitalDepartments)
    cur.executemany("INSERT INTO doctors VALUES(null, ?, ?, ?, ?, ?, ?, ?);", dataDoctors)
    cur.executemany("INSERT INTO diagnos VALUES(null, ?, ?, ?);", dataDiagnoz)
    cur.executemany("INSERT INTO patients_diagnos VALUES(?, ?);", dataManyToMany)
    cur.executemany("INSERT INTO patients VALUES(null, ?, ?, ?, ?, ?);", dataPatient)
    conn.commit()




if __name__ == "__main__":
    #createDB()
    #inputData()

    cur.execute("SELECT doctors.fulName, doctors.sex, doctors.birthDate, doctors.Salary  FROM doctors where DATE('now') - doctors.birthDate >= 35;")
    printFormatRecod(cur.fetchall(), ("ПІБ  лікаря", "Стать", "Дата народження", "Заробітна плата"))

    '''
    #
    #cur.execute("SELECT fulName, jobTitle, sex, birthDate, Salary  FROM doctors;")
    #printFormatRecod(cur.fetchall(), ("ПІБ  лікаря", "Стать", "Дата народження", "Заробітна плата"))
    cur.execute("UPDATE doctors SET fulName  = 'ОЛІЙНИК АЛІНА ДМИТРІВНА' WHERE  fulName = 'КРАВЧЕНКО АЛІНА ДМИТРІВНА';")
    cur.execute("SELECT fulName, jobTitle, sex, birthDate, Salary  FROM doctors WHERE sex = 'жінка' and fulName LIKE '%ОЛІЙНИК%';")
    printFormatRecod(cur.fetchall(), ("ПІБ  лікаря", "Стать", "Дата народження", "Заробітна плата"))
    
    

    cur.execute("DELETE FROM patients WHERE fulName ='ШЕВЧЕНКО ЛЮДМИЛА МИХАЙЛІВНА';")
    cur.execute("""INSERT INTO patients (fulName , sex , birthDate , receiptDate, doctorID  )
    VALUES('ТКАЧЕНКО ИННА БОРИСІВНА',  'жінка', '1960-12-12', '2022-12-25', 4);""")
    cur.execute("SELECT fulName, sex, birthDate, receiptDate    FROM patients WHERE sex = 'жінка';")
    printFormatRecod(cur.fetchall(), ("ПІБ  пацієнта", "Стать", "Дата народження", "Дата вступу"))
    '''



    cur.execute("""
    SELECT hospitalDepartments.deparName, doctors.fulName, patients.fulName, patients.sex, patients.birthDate 
    FROM  hospitalDepartments  LEFT  JOIN doctors  ON 
    doctors.ID  = hospitalDepartments.ID INNER
    JOIN patients  ON patients.doctorID  = doctors.ID WHERE  patients.sex LIKE 'чоловік';
    """)
    printFormatRecod(cur.fetchall(), ("Відділення", "ПІБ лікуючого лікаря", "ПІБ пацієнта", "Стать пацієнта", "Дата народження пацієнт"))


    getIngexTemp = cur.execute("SELECT ID FROM patients WHERE fulName='КОВАЛЬЧУК МИКОЛА ІВАНОВИЧ'").fetchone()[0]
    cur.execute("SELECT name FROM diagnos d  LEFT JOIN patients_diagnos gp ON gp.diagnos_id = d.id  WHERE gp.patients_id = "+str(getIngexTemp)+"");
    printFormatRecod(cur.fetchall(), ("КОВАЛЬЧУК МИКОЛА ІВАНОВИЧ має такі захварювання:",))
    getIngexTemp = cur.execute("SELECT ID FROM diagnos WHERE name='кардіосклероз'").fetchone()[0]
    cur.execute("SELECT fulName FROM patients p  LEFT JOIN patients_diagnos gp ON gp.patients_id = p.id  WHERE gp.diagnos_id = "+str(getIngexTemp));
    printFormatRecod(cur.fetchall(), ("кардіосклероз наявний у наступних пациентів:",))



    cur.execute("SELECT patients.sex, COUNT(patients.sex) FROM patients GROUP BY patients.sex;")
    printFormatRecod(cur.fetchall(), ("Стать:", "Кількість докторів"))

    cur.execute("""
    SELECT hospitalDepartments.deparName, SUM(doctors.Salary)  FROM  hospitalDepartments  INNER 
    JOIN doctors  ON doctors.depertID  = hospitalDepartments.ID GROUP BY hospitalDepartments.deparName;
    """)
    printFormatRecod(cur.fetchall(), ("Назва відділення:", "Видатки на утримання (грн)"))


    cur.execute("SELECT SUM(doctors.Salary) FROM doctors;")
    printFormatRecod(cur.fetchall(), ("Загальна сумма посадових окладів лікарні (грн):",))


    #Complex request
    cur.execute("SELECT fullDescrp  FROM hospitalDepartments WHERE deparName LIKE '%терапев%';")
    one_result = cur.fetchall()
    decoded = base64.b64decode(one_result[0][0])
    myDocuments = os.path.join(os.path.expanduser('~'),'Downloads')
    f = open(myDocuments+"\pologennia.docx", "wb")
    print("Document saved by the next path: "+myDocuments+"\pologennia.docx")
    f.write(decoded)
    f.close()

    cur.execute("""
    SELECT deparName, doctors.fulName, doctors.jobTitle, doctors.sex, doctors.birthDate, doctors.Salary, doctors.photo   FROM  hospitalDepartments  INNER 
    JOIN doctors  ON doctors.depertID  = hospitalDepartments.ID  WHERE doctors.fulName = 'БОНДАРЕНКО МИКОЛА ІВАНОВИЧ';
    """)
    getDocResume(cur.fetchone())
    conn.commit()
