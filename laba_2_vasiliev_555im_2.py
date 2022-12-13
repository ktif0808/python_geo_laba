import matplotlib.pyplot as plt
import numpy as np
import mplcursors
import matplotlib.dates as mdates
from datetime import datetime
from geopy import distance
from pyproj import Geod
from geographiclib.geodesic import Geodesic
import folium
import sys
import docx
import os


class makeTableDocx:
    @staticmethod
    def createTab1(arrLat, arrLong, arrDis, arrTime, arrSpeed):
        converArrTime = []
        print(arrTime)
        for i in range(len(arrTime)):
            converArrTime.append(int((arrTime[i] - arrTime[0]).total_seconds()))
        doc = docx.Document()
        doc.add_heading('Таблиця 1', 0)
        records = []
        for i in range(len(arrLat)):
            records.append([str(i+1), str(arrLat[i]) , str(arrLong[i]) , str(arrDis[i]), str(converArrTime[i]), str(arrSpeed[i])])

        menuTable = doc.add_table(rows=1, cols=6)
        menuTable.style='Table Grid'
        hdr_Cells = menuTable.rows[0].cells
        hdr_Cells[0].text = "Номер"
        hdr_Cells[1].text = 'Широта'
        hdr_Cells[2].text = 'Довгота'
        hdr_Cells[3].text = 'Відстань(м)'
        hdr_Cells[4].text = 'Час польоту'
        hdr_Cells[5].text = 'швидкість польоту (Км/год)'
        for ID, lat, long, dis, time, speed in records:
            row_Cells = menuTable.add_row().cells
            row_Cells[0].text = str(ID)
            row_Cells[1].text = lat
            row_Cells[2].text = long
            row_Cells[3].text = dis
            row_Cells[4].text = time
            row_Cells[5].text = speed
        doc.save("\\python_laba\\resourse\\table.docx")
        #os.system("start table.docx")

    @staticmethod
    def createTab2(sumDis, averSpeed, maxAltitude, minAltitude, maxSpeed, minSpeed, timeOfFlight):
        doc = docx.Document()
        doc.add_heading('Таблиця 2', 0)
        records = [[1, 'Сумарна довжина маршруту польоту ', str(sumDis), 'м'],
                   [2, 'Середня швидкість польоту', str(averSpeed), 'Км\г'],
                   [3, 'Макс. Висота польоту', str(maxAltitude), 'м'],
                   [4, 'Мін. Висота польоту', str(minAltitude), 'м'],
                   [5, 'Макс. швидкість польоту', str(maxSpeed), 'Км\г'],
                   [6, 'Мін. швидкість польоту', str(minSpeed), 'Км\г'],
                   [7, 'Загальний час польоту', str(timeOfFlight), 'с']
                   ]

        menuTable = doc.add_table(rows=1, cols=4)
        menuTable.style='Table Grid'
        hdr_Cells = menuTable.rows[0].cells
        hdr_Cells[0].text = "Номер"
        hdr_Cells[1].text = 'Параметр'
        hdr_Cells[2].text = 'Значення'
        hdr_Cells[3].text = 'Одиниця виміру'

        for ID, cell1, cell2, cell3 in records:
            row_Cells = menuTable.add_row().cells
            row_Cells[0].text = str(ID)
            row_Cells[1].text = cell1
            row_Cells[2].text = cell2
            row_Cells[3].text = cell3

        doc.save("\\python_laba\\resourse\\table2.docx")
        #os.system("start table2.docx")




class DistanceCalculator:

    #I will use this method for all my examples as a default
    @staticmethod
    def getDistanceGeopy(point1, point2):
        distance_2d= distance.distance(point1[:2], point2[:2]).m
        distance_3d = np.sqrt(distance_2d**2 + (point1[2] - point2[2])**2)
        return np.around([distance_3d], decimals=2, out=None)[0]

    @staticmethod
    def getDistanceGeod(point1, point2):
        g = Geod(ellps='WGS84')
        azimuth1, azimuth2, distance_2d = g.inv(point1[1], point1[0], point2[1], point2[0])
        distance_3d = np.hypot(distance_2d,point2[2]-point1[2])
        return np.around([distance_3d], decimals=2, out=None)[0]

    @staticmethod
    def getDistanceGeodesic(point1, point2):
        geod = Geodesic.WGS84
        g = geod.Inverse(point1[0], point1[1],point2[0],point2[1])
        distance_3d = np.hypot(g['s12'],point2[2]-point1[2])
        return np.around([distance_3d], decimals=2, out=None)[0]

#sys.exit(0)

class DrawPlot:

    @staticmethod
    def buildPlot2D(xData, yData, xlabel, ylabel, title, colorPlot):
        def textRun(annotation):
            syt = annotation.replace('x', xlabel)
            syt = syt.replace('y', ylabel)
            return syt
        plt.style.use('seaborn-whitegrid')
        myFmt = mdates.DateFormatter('%H:%M:%S')
        fig, ax = plt.subplots()

        faceColor = None
        pointColor = None
        if colorPlot == "blue":
            faceColor = 'lightblue'
            pointColor = 'darkred'
        elif colorPlot == "darkgreen":
            faceColor = 'antiquewhite'
            pointColor = 'red'
        elif colorPlot == 'darkorange':
            faceColor = 'lavender'
            pointColor = 'darkred'
        else:
            raise Exception("Incorect input")

        fig.set(facecolor = faceColor)
        ax.plot(xData, yData, color=colorPlot)
        ax.set_xlabel(xlabel, fontsize=20)
        ax.set_ylabel(ylabel, fontsize=20)
        plt.title(title, fontsize=20)
        dots = ax.scatter(xData, yData, color=pointColor)
        plt.legend(['dependency', 'checkpoint'], loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=12, frameon=True)
        ax.xaxis.set_major_formatter(myFmt)
        crs = mplcursors.cursor(dots, hover=True)
        crs.connect("add", lambda sel: sel.annotation.set_text(textRun(sel.annotation.get_text())))
        plt.show()


    @staticmethod
    def buildPlot3D(xData, yData, zData, xlabel, ylabel, zLabel, title):
        print(xData)
        print(yData)
        print(zData)
        fig = plt.figure()
        fig.set(facecolor = 'whitesmoke')
        ax = fig.add_subplot(projection='3d')
        plt.title(title)

        ax.plot3D(xData, yData, zData, 'red', label="flight path")
        ax.scatter3D(xData, yData, zData, label="Check point")

        ax.set_xlabel(xlabel, fontsize=15)
        ax.set_ylabel(ylabel, fontsize=15)
        ax.set_zlabel(zLabel, fontsize=15)
        ax.legend(loc=2)

        iter = 1
        for x,y,z in zip(xData, yData, zData):
            if iter==1:
                ax.text(x, y, z, 'start', horizontalalignment='left', size='medium', color='black')
            elif iter==20:
                ax.text(x, y, z, 'end', horizontalalignment='left', size='medium', color='black')
            iter+=1
        plt.show()

    @staticmethod
    def buildFoliumMap(xData, yData):
        print(xData)
        print(yData)
        m = folium.Map([xData[0], yData[0]], zoom_start=9, control_scale = True)
        for i in range(1, len(xData)):
            folium.PolyLine(locations = [[xData[i-1] , yData[i-1]], [xData[i], yData[i] ]],
                        line_opacity = 0.5, color='red').add_to(m)

        icon1 = folium.features.CustomIcon(icon_image="\\python_laba\\resourse\\start.png", icon_size="40")
        folium.Marker(location=[xData[0], yData[0]],icon=icon1).add_to(m)

        icon2 = folium.features.CustomIcon(icon_image="\\python_laba\\resourse\\finish.png", icon_size="40")
        folium.Marker(location=[xData[len(xData)-1], yData[len(yData)-1]],icon=icon2).add_to(m)

        m.save("\\python_laba\\resourse\\mapFolium.html")



    @staticmethod
    def drawPlot(plan, dependency):
        if dependency=="time-speed":
            DrawPlot.buildPlot2D(plan._timeHigh, plan._arrSpeed, 'check time (control point)', \
                                 'speed (m/s)', 'The plot of dependency of speed to time', 'blue')
        elif dependency=="time-altitude":
            DrawPlot.buildPlot2D(plan._timeHigh, plan._arrHight, 'check time (control point)', \
                                 'altitude (m)', 'The plot of dependency of altitude to time', 'darkgreen')
        elif dependency=="time-distance":
            DrawPlot.buildPlot2D(plan._timeHigh, plan._arrDis, 'check time (control point)', \
                                 'distance (m)', 'The plot of dependency of distance to time', 'darkorange')
        elif dependency=="flight-path3D":
            DrawPlot.buildPlot3D(plan._arrLong, plan._arrLat, plan._arrHight, 'Longitude (m)', \
                                 'Latitude (m)', 'altitude (m)','The plot of 3D path of the airplane')
        elif dependency=="floliumMap":
            DrawPlot.buildFoliumMap(plan._arrCorLat, plan._arrCorLong)
        else:
            raise Exception("Incorect input")


class Airplane:

    def __init__(self, arrData, distanceMethod="Geopy"):
        self._arrData = arrData
        self.executeFlight(self._arrData, distanceMethod)

    def executeFlight(self, arrData, distanceMethod="Geopy"):

        sumDis = 0
        maxSpeed = 0
        minSpeed = 99999999
        minFlight = float(arrData[0][9])
        maxFlight = float(arrData[0][9])
        arrDis = [0]
        arrSpeed = [0]
        arrHight = [arrData[0][9]]

        #for 3d
        arrLat = [0]
        arrLong = [0]

        #for folium map
        arrCorLat = [float(arrData[0][2])/100]
        arrCorLong = [float(arrData[0][4])/100]



        for i in range(1, len(arrData)):
            point1 = [float(arrData[i-1][2])/100, float(arrData[i-1][4])/100, float(arrData[i-1][9])]
            point2 = [float(arrData[i][2])/100, float(arrData[i][4])/100, float(arrData[i][9])]
            dis2point = None
            if distanceMethod=="Geopy":
                dis2point = DistanceCalculator.getDistanceGeopy(point1,point2)
            elif distanceMethod=="Geod":
                dis2point = DistanceCalculator.getDistanceGeod(point1,point2)
            elif distanceMethod=="Geodesic":
                dis2point = DistanceCalculator.getDistanceGeodesic(point1,point2)
            else:
                raise Exception("Incorect input")
            sumDis+=dis2point
            arrDis.append(sumDis)
            arrDis[i] = np.around([float(arrDis[i])], decimals=2, out=None)[0]
            arrSpeed.append(float(dis2point))
            arrHight.append(float(arrData[i][9]))
            maxSpeed = max(dis2point, maxSpeed)
            minSpeed = min(dis2point, minSpeed)
            minFlight = min(float(arrData[i][9]), minFlight)
            maxFlight = max(float(arrData[i][9]), maxFlight)


            #for 3d flight
            vlong = DistanceCalculator.getDistanceGeopy([0.0, float(arrData[i-1][4])/100, 0.0], [0.0, float(arrData[i][4])/100, 0.0])
            vlat = DistanceCalculator.getDistanceGeopy([float(arrData[i-1][2])/100, 0.0, 0.0], [float(arrData[i][2])/100, 0.0, 0.0])
            arrLong.append(arrLong[i-1]+vlong)
            arrLat.append(arrLat[i-1]+vlat)
            arrLong[i] = np.around([float(arrLong[i])], decimals=2, out=None)[0]
            arrLat[i] = np.around([float(arrLat[i])], decimals=2, out=None)[0]

            #for folium map
            arrCorLat.append(float(arrData[i][2])/100)
            arrCorLong.append(float(arrData[i][4])/100)
            arrCorLat[i] = np.around([float(arrCorLat[i])], decimals=6, out=None)[0]
            arrCorLong[i] = np.around([float(arrCorLong[i])], decimals=6, out=None)[0]

        #convert to numpy array
        arrSpeed = np.array(arrSpeed)
        arrSpeed = arrSpeed.astype(np.float64)
        arrHight = np.array(arrHight)
        arrHight = arrHight.astype(np.float64)
        arrDis = np.array(arrDis)
        arrDis = arrDis.astype(np.float64)

        arrLong = np.array(arrLong)
        arrLong = arrLong.astype(np.float64)
        arrLat = np.array(arrLat)
        arrLat = arrLat.astype(np.float64)


        #time plotting
        timeHigh = []
        for i in range(0, len(arrData)):
            strD = arrData[i][1][:-3]
            datetime_object = datetime.strptime(strD, '%H%M%S')
            timeHigh.append(datetime_object)
        #total seconds
        resultTime = timeHigh[len(timeHigh)-1] - timeHigh[0]



        #data request
        self._maxSpeed = maxSpeed
        self._minSpeed = minSpeed
        self._minFlight = minFlight
        self._maxFlight = maxFlight
        self._sumDis = sumDis
        self._averSpeed =sumDis/(len(arrDis)-1)
        self._timeFlight = resultTime.total_seconds()

        #array for plotting
        self._timeHigh = timeHigh
        self._arrSpeed = arrSpeed
        self._arrHight = arrHight
        self._arrDis = arrDis


        #for 3d plotting
        self._arrLong = arrLong
        self._arrLat = arrLat

        #for folium plotting
        self._arrCorLong = arrCorLong
        self._arrCorLat = arrCorLat

    def convertMCtoKG(self, value):
        value = value/1000*3600
        return value

    def buildTableDoc1(self):
        makeTableDocx.createTab1(self._arrCorLat, self._arrCorLong, self._arrDis, self._timeHigh, self._arrSpeed)
    def buildTableDoc2(self):
        makeTableDocx.createTab2(self.getTotalDistance(), self.getAverSpeed(), self._maxFlight, self._minFlight, self._maxSpeed, self._minSpeed, self._timeFlight)

    #get requered paraments to our PZ
    def getMinAltitude(self):
        return np.around([self._minFlight], decimals=2, out=None)[0]
    def getMaxAltitude(self):
        return np.around([self._maxFlight], decimals=2, out=None)[0]
    def getTotalDistance(self):
        return np.around([self._sumDis], decimals=2, out=None)[0]
    def getTimeFlight(self):
        return np.around([self._timeFlight], decimals=2, out=None)[0]
    def getAverSpeed(self, convert=False):
        if convert==False:
            return np.around([self._averSpeed], decimals=2, out=None)[0]
        else:
            return np.around([self.convertMCtoKG(self._averSpeed)], decimals=2, out=None)[0]
    def getMaxSpeed(self, convert=False):
        if convert==False:
            return np.around([self._maxSpeed], decimals=2, out=None)[0]
        else:
            return np.around([self.convertMCtoKG(self._maxSpeed)], decimals=2, out=None)[0]
    def getMinSpeed(self, convert=False):
        if convert==False:
            return np.around([self._minSpeed], decimals=2, out=None)[0]
        else:
            return np.around([self.convertMCtoKG(self._minSpeed)], decimals=2, out=None)[0]


    def buildDepenency(self, dependency):
        DrawPlot.drawPlot(self, dependency)

    def setData(self, arrData, distanceMethod="Geopy"):
        self._arrData = arrData
        self.executeFlight(self._arrData, distanceMethod)
    def getData(self):
        return self._arrData
    def printData(self):
        for i in range(len(self._arrData)):
            print(i+1, self._arrData[i])



def readFile(pathFile):
    file1 = open(pathFile, 'r')
    Lines = file1.readlines()
    arrData = []
    for line in Lines:
        st = list(map(str, line.split(',')))
        arrData.append(st)
    return arrData


if __name__ == "__main__":
    dataFlight = readFile('\\python_laba\\resourse\\data.csv')
    airplane1 = Airplane(dataFlight)
    airplane1.buildTableDoc1()
    airplane1.buildTableDoc2()
    #sys.exit(0)
    #plan1.printData()
    print("The distance that the airplane has flight:", airplane1.getTotalDistance(), "m")
    print("Maximum speed of the airplane:", airplane1.getMaxSpeed(),"m/c")
    print("Maximum speed of the airplane:", airplane1.getMaxSpeed(convert=True), "km/h")
    print("Minimum speed of the airplane:", airplane1.getMinSpeed(),"m/c")
    print("Minimum speed of the airplane:", airplane1.getMinSpeed(convert=True), "km/h")
    print("Average speed of the airplane:", airplane1.getAverSpeed(),"m/c")
    print("Average speed of the airplane:", airplane1.getAverSpeed(convert=True), "km/h")
    print("Maximum altitude of the airplane:", airplane1.getMaxAltitude(), "m")
    print("Minimum altitude of the airplane:", airplane1.getMinAltitude(), "m")
    print("Total time of the flight:", airplane1.getTimeFlight(), "s")

    airplane1.buildDepenency(dependency="time-altitude")
    airplane1.buildDepenency(dependency="time-speed")
    airplane1.buildDepenency(dependency="time-distance")
    airplane1.buildDepenency(dependency="flight-path3D")
    airplane1.buildDepenency(dependency="floliumMap")

